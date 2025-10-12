from mcp.server.fastmcp import FastMCP
from typing import Optional, Union, List, Dict
from pydantic import BaseModel
import os

# 导入 docx 相关库
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

mcp = FastMCP('word-mcp')

# 全局变量存储当前活动的 Word 文档
ACTIVE_DOCUMENT: Optional[Document] = None
ACTIVE_DOCUMENT_PATH: Optional[str] = None

# 根目录配置
DEFAULT_WORD_MCP_ROOT = '.word-mcp'
_env_root = os.environ.get('WORD_MCP_ROOT')
if _env_root and os.path.isabs(_env_root):
    WORD_MCP_ROOT = _env_root
else:
    WORD_MCP_ROOT = DEFAULT_WORD_MCP_ROOT

def ensure_root_dir():
    """确保根目录存在"""
    try:
        os.makedirs(WORD_MCP_ROOT, exist_ok=True)
    except Exception as e:
        print(f"创建根目录时出错: {e}")
ensure_root_dir()

def get_document_path(name: str) -> str:
    """根据文档名称返回文档存储路径"""
    return os.path.join(WORD_MCP_ROOT, f"{name}.docx")

class WordResult(BaseModel):
    success: bool
    message: str
    output: Optional[Union[str, int, List[str], Dict]] = None

@mcp.prompt()
def word_operations_prompt():
    """指导 AI 使用 Word 文档操作工具"""
    return """
你是一个擅长使用 Word 文档进行文档创建和编辑的助手。你可以帮助用户创建、编辑和格式化 Word 文档。

可用的操作包括：
1. 创建新文档
2. 打开现有文档
3. 添加段落、标题和文本
4. 格式化文本（字体、大小、颜色、对齐方式）
5. 添加表格并设置表格内容
6. 添加图片
7. 保存文档

请根据用户的需求选择合适的工具进行操作。在开始操作前，通常需要先创建一个新文档或打开一个现有文档。
"""

@mcp.tool()
def word_create_document(name: str) -> WordResult:
    """
    创建一个新的 Word 文档
    `name`: 文档名称（不需要包含.docx扩展名）
    """
    global ACTIVE_DOCUMENT, ACTIVE_DOCUMENT_PATH
    
    try:
        # 创建新文档
        ACTIVE_DOCUMENT = Document()
        ACTIVE_DOCUMENT_PATH = get_document_path(name)
        
        return WordResult(
            success=True, 
            message=f"文档 '{name}' 创建成功", 
            output={"name": name, "path": ACTIVE_DOCUMENT_PATH}
        )
    except Exception as e:
        return WordResult(
            success=False, 
            message=f"创建文档失败: {str(e)}"
        )

@mcp.tool()
def word_open_document(name: str) -> WordResult:
    """
    打开一个现有的 Word 文档
    `name`: 文档名称（不需要包含.docx扩展名）
    """
    global ACTIVE_DOCUMENT, ACTIVE_DOCUMENT_PATH
    
    try:
        document_path = get_document_path(name)
        
        if not os.path.exists(document_path):
            return WordResult(
                success=False, 
                message=f"文档 '{name}' 不存在"
            )
        
        # 打开文档
        ACTIVE_DOCUMENT = Document(document_path)
        ACTIVE_DOCUMENT_PATH = document_path
        
        # 获取文档基本信息
        paragraph_count = len(ACTIVE_DOCUMENT.paragraphs)
        table_count = len(ACTIVE_DOCUMENT.tables)
        
        return WordResult(
            success=True, 
            message=f"文档 '{name}' 打开成功", 
            output={
                "name": name, 
                "path": document_path,
                "paragraphs": paragraph_count,
                "tables": table_count
            }
        )
    except Exception as e:
        return WordResult(
            success=False, 
            message=f"打开文档失败: {str(e)}"
        )

@mcp.tool()
def word_add_heading(text: str, level: int = 1) -> WordResult:
    """
    添加一个标题
    `text`: 标题文本
    `level`: 标题级别 (1-9, 1是最高级别)
    """
    global ACTIVE_DOCUMENT
    
    if not ACTIVE_DOCUMENT:
        return WordResult(
            success=False, 
            message="没有活动的文档。请先创建或打开一个文档"
        )
    
    try:
        # 添加标题
        heading = ACTIVE_DOCUMENT.add_heading(text, level=level)
        
        return WordResult(
            success=True, 
            message=f"添加 {level} 级标题成功", 
            output={"text": text, "level": level}
        )
    except Exception as e:
        return WordResult(
            success=False, 
            message=f"添加标题失败: {str(e)}"
        )

@mcp.tool()
def word_add_paragraph(text: str, style: Optional[str] = None) -> WordResult:
    """
    添加一个段落
    `text`: 段落文本
    `style`: 段落样式 (可选)
    """
    global ACTIVE_DOCUMENT
    
    if not ACTIVE_DOCUMENT:
        return WordResult(
            success=False, 
            message="没有活动的文档。请先创建或打开一个文档"
        )
    
    try:
        # 添加段落
        paragraph = ACTIVE_DOCUMENT.add_paragraph(text, style=style)
        
        return WordResult(
            success=True, 
            message="添加段落成功", 
            output={"text": text, "style": style}
        )
    except Exception as e:
        return WordResult(
            success=False, 
            message=f"添加段落失败: {str(e)}"
        )

@mcp.tool()
def word_add_table(rows: int, cols: int, data: Optional[List[List[str]]] = None) -> WordResult:
    """
    添加一个表格
    `rows`: 行数
    `cols`: 列数
    `data`: 表格数据 (可选，二维列表)
    """
    global ACTIVE_DOCUMENT
    
    if not ACTIVE_DOCUMENT:
        return WordResult(
            success=False, 
            message="没有活动的文档。请先创建或打开一个文档"
        )
    
    try:
        # 添加表格
        table = ACTIVE_DOCUMENT.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # 如果有数据，填充表格
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_data in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_data)
        
        return WordResult(
            success=True, 
            message="添加表格成功", 
            output={"rows": rows, "cols": cols, "data": data}
        )
    except Exception as e:
        return WordResult(
            success=False, 
            message=f"添加表格失败: {str(e)}"
        )

@mcp.tool()
def word_format_text(
    paragraph_index: int, 
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_size: Optional[int] = None,
    font_name: Optional[str] = None,
    color: Optional[str] = None
) -> WordResult:
    """
    格式化文本
    `paragraph_index`: 段落索引
    `bold`: 是否加粗
    `italic`: 是否斜体
    `underline`: 是否下划线
    `font_size`: 字体大小
    `font_name`: 字体名称
    `color`: 字体颜色 (十六进制格式，如 #FF0000)
    """
    global ACTIVE_DOCUMENT
    
    if not ACTIVE_DOCUMENT:
        return WordResult(
            success=False, 
            message="没有活动的文档。请先创建或打开一个文档"
        )
    
    try:
        if paragraph_index >= len(ACTIVE_DOCUMENT.paragraphs):
            return WordResult(
                success=False, 
                message=f"段落索引 {paragraph_index} 超出范围"
            )
        
        paragraph = ACTIVE_DOCUMENT.paragraphs[paragraph_index]
        
        # 应用格式
        for run in paragraph.runs:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if underline is not None:
                run.underline = underline
            if font_size is not None:
                run.font.size = Pt(font_size)
            if font_name is not None:
                run.font.name = font_name
            if color is not None:
                # 将十六进制颜色转换为RGB
                color = color.lstrip('#')
                rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
                run.font.color.rgb = RGBColor(*rgb)
        
        return WordResult(
            success=True, 
            message="格式化文本成功", 
            output={"paragraph_index": paragraph_index}
        )
    except Exception as e:
        return WordResult(
            success=False, 
            message=f"格式化文本失败: {str(e)}"
        )

# 

@mcp.tool()
def word_set_alignment(paragraph_index: int, alignment: str) -> WordResult:
    """
    设置段落对齐方式
    `paragraph_index`: 段落索引
    `alignment`: 对齐方式 (left, center, right, justify)
    """
    global ACTIVE_DOCUMENT
    
    if not ACTIVE_DOCUMENT:
        return WordResult(
            success=False, 
            message="没有活动的文档。请先创建或打开一个文档"
        )
    
    try:
        if paragraph_index >= len(ACTIVE_DOCUMENT.paragraphs):
            return WordResult(
                success=False, 
                message=f"段落索引 {paragraph_index} 超出范围"
            )
        
        paragraph = ACTIVE_DOCUMENT.paragraphs[paragraph_index]
        
        # 设置对齐方式
        if alignment == "left":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "justify":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            return WordResult(
                success=False, 
                message=f"不支持的对齐方式: {alignment}"
            )
        
        return WordResult(
            success=True, 
            message=f"设置对齐方式成功: {alignment}", 
            output={"paragraph_index": paragraph_index, "alignment": alignment}
        )
    except Exception as e:
        return WordResult(
            success=False, 
            message=f"设置对齐方式失败: {str(e)}"
        )

@mcp.tool()
def word_save_document() -> WordResult:
    """保存当前活动的文档"""
    global ACTIVE_DOCUMENT, ACTIVE_DOCUMENT_PATH
    
    if not ACTIVE_DOCUMENT or not ACTIVE_DOCUMENT_PATH:
        return WordResult(
            success=False, 
            message="没有活动的文档可保存"
        )
    
    try:
        ACTIVE_DOCUMENT.save(ACTIVE_DOCUMENT_PATH)
        
        # 获取文档统计信息
        paragraph_count = len(ACTIVE_DOCUMENT.paragraphs)
        table_count = len(ACTIVE_DOCUMENT.tables)
        
        return WordResult(
            success=True, 
            message=f"文档已保存到: {ACTIVE_DOCUMENT_PATH}", 
            output={
                "path": ACTIVE_DOCUMENT_PATH,
                "paragraphs": paragraph_count,
                "tables": table_count
            }
        )
    except Exception as e:
        return WordResult(
            success=False, 
            message=f"保存文档失败: {str(e)}"
        )

@mcp.tool()
def word_save_as_document(name: str) -> WordResult:
    """
    将当前文档另存为
    `name`: 新文档名称
    """
    global ACTIVE_DOCUMENT
    
    if not ACTIVE_DOCUMENT:
        return WordResult(
            success=False, 
            message="没有活动的文档可保存"
        )
    
    try:
        new_path = get_document_path(name)
        ACTIVE_DOCUMENT.save(new_path)
        
        # 更新活动文档路径
        global ACTIVE_DOCUMENT_PATH
        ACTIVE_DOCUMENT_PATH = new_path
        
        # 获取文档统计信息
        paragraph_count = len(ACTIVE_DOCUMENT.paragraphs)
        table_count = len(ACTIVE_DOCUMENT.tables)
        
        return WordResult(
            success=True, 
            message=f"文档已另存为: {new_path}", 
            output={
                "path": new_path,
                "paragraphs": paragraph_count,
                "tables": table_count
            }
        )
    except Exception as e:
        return WordResult(
            success=False, 
            message=f"另存为文档失败: {str(e)}"
        )

@mcp.tool()
def word_get_document_info() -> WordResult:
    """获取当前文档的信息"""
    global ACTIVE_DOCUMENT
    
    if not ACTIVE_DOCUMENT:
        return WordResult(
            success=False, 
            message="没有活动的文档"
        )
    
    try:
        # 获取文档统计信息
        paragraph_count = len(ACTIVE_DOCUMENT.paragraphs)
        table_count = len(ACTIVE_DOCUMENT.tables)
        
        # 获取段落预览
        paragraphs_preview = []
        for i, paragraph in enumerate(ACTIVE_DOCUMENT.paragraphs[:5]):  # 只显示前5个段落
            if paragraph.text.strip():  # 只显示非空段落
                paragraphs_preview.append({
                    "index": i,
                    "text": paragraph.text[:100] + "..." if len(paragraph.text) > 100 else paragraph.text
                })
        
        return WordResult(
            success=True, 
            message="文档信息获取成功", 
            output={
                "path": ACTIVE_DOCUMENT_PATH,
                "paragraphs": paragraph_count,
                "tables": table_count,
                "preview": paragraphs_preview
            }
        )
    except Exception as e:
        return WordResult(
            success=False, 
            message=f"获取文档信息失败: {str(e)}"
        )